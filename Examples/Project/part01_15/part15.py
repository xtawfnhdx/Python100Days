from PIL import Image, ImageFilter


# 显示图片
def Show():
    image = Image.open('E:/My/Pic/桌面/IMG_0234.JPG')
    image.format, image.size, image.mode
    ('JPEG', (500, 750), 'RGB')
    image.show()


# 剪切图片
def Crop():
    image = Image.open('E:/My/Pic/桌面/IMG_0234.JPG')
    rect = 80, 20, 310, 360
    image.crop(rect).show()


# 缩略图
def Thumbnail():
    image = Image.open('E:/My/Pic/桌面/IMG_0234.JPG')
    size = 128, 128
    image.thumbnail(size)
    image.show()


# 剪切粘贴图片
def Paste():
    image1 = Image.open('E:/My/Pic/桌面/IMG_0234.JPG')
    image2 = Image.open('E:/My/Pic/桌面/IMG_4684.JPG')
    rect = 2000, 1500, 4000, 3000
    guido_head = image2.crop(rect)
    width, height = guido_head.size
    image1.paste(guido_head, (2000, 1500))
    image1.show()


# 图像翻转
def Rotate():
    image = Image.open('E:/My/Pic/桌面/IMG_0234.JPG')
    image.rotate(40).show()
    image.transpose(Image.FLIP_LEFT_RIGHT).show()


# 像素操作
def Putpixel():
    image = Image.open('E:/My/Pic/桌面/IMG_0234.JPG')
    for x in range(3000, 4500):
        for y in range(2000, 3000):
            image.putpixel((x, y), (128, 128, 128))
    image.show()


# 滤镜效果
def Filter():
    image = Image.open('E:/My/Pic/重要/IMG_4378.JPG')
    # image.filter(ImageFilter.CONTOUR).show()
    # image.filter(ImageFilter.EMBOSS).show()
    image.filter(ImageFilter.FIND_EDGES).show()


import datetime
from openpyxl import Workbook


# 操作Excel
def WorkbookDef():
    wb = Workbook()
    ws = wb.active
    ws['A1'] = 42
    ws.append([1, 2, 3])
    ws.append([4, 5, 6])
    ws['A2'] = datetime.datetime.now()
    wb.save('part15.xlsx')


from docx import Document
from docx.shared import Inches


def WordDef():
    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading,level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    document.add_paragraph(
        'second item in ordered list', style='List Number'
    )

    document.add_picture('E:/My/Pic/桌面/IMG_0234.JPG', width=Inches(3))
    records = (
        (3, '101', 'Span'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam,spam,eggs,and span')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    # 增加一个新的页面
    document.add_page_break()
    document.save('part15.docx')


if __name__ == "__main__":
    # Show()
    # Crop()
    # Thumbnail()
    # Paste()
    # Rotate()
    # Putpixel()
    # Filter()
    # WorkbookDef()
    WordDef()